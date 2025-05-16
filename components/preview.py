import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from io import BytesIO

# Hardcoded JSON data (as provided)
resume_data = {
    "creation_date": "2025-03-13",
    "personal_info": {
        "full_name": "Onkabetse Bernice",
        "surname": "Seboko",
        "birth_date": "1972-12-02",
        "sa_citi": "Yes",
        "work_permit": "No",
        "id_num": "7212021097088",
        "pas_num": "",
        "race": "African",
        "nationality": "Republic of South Africa",
        "gender": "Female",
        "disability": "No",
        "p_location": "Mahikeng",
        "phone": "0765013356",
        "email": "sebokoonkabetse9@gmail.com",
        "linkedin": "",
        "other_profiles": ""
    },
    "personal_add_info": {},
    "corr_pref": {
        "pub_ser": "Yes",
        "pre_depart": "Mahikeng  Local Municipality. (PR) Ward councillor"
    },
    "current_qual": {
        "current_qual": ""
    },
    "professional_summary": {},
    "work_experience": [
        {
            "job_title": "ABET TEACHER",
            "company": "SUNLI PROJECT",
            "dates": "2002 - 2003",
            "reason": "End of contract."
        },
        {
            "job_title": "PR Councillor",
            "company": "Mahikeng Local Municipality",
            "dates": "2006 - 2011",
            "reason": "End of contract."
        },
        {
            "job_title": "Registration Operator",
            "company": "CPS",
            "dates": "2012 ( 3 Months)",
            "reason": "End of contract."
        },
        {
            "job_title": "Child Cartetaker",
            "company": "STADT PRIMARY SCHOOL",
            "dates": "2016 - 2021",
            "reason": "End of contract."
        },
        {
            "job_title": "Data Capturer",
            "company": "Magogoe Clinic",
            "dates": "2025 - Present",
            "reason": "Career growth."
        }
    ],
    "education": [
        {
            "degree": "Certificate in ABET",
            "field": "Education",
            "institution": "University of South Africa",
            "location": "Mahikeng",
            "graduation_date": "2002"
        },
        {
            "degree": "DIPLOMA IN DATA PROCESSING",
            "field": "Information Systems",
            "institution": "Dynamic Comm. College",
            "location": "Mahikeng",
            "graduation_date": "1995"
        },
        {
            "degree": "Matric ",
            "field": "Matric",
            "institution": "Leteane High School",
            "location": "Mahikeng",
            "graduation_date": "1994"
        }
    ],
    "languages": [],
    "skills": {
        "technical": "Information Systems\nMicrosoft\nSpreadsheet\nAccounting \nBookkeeping",
        "software": "Microsoft 365",
        "methodologies": "",
        "languages": "Setswana\nEnglish\nSesotho",
        "soft_skills": "Communication\nCollaboration"
    },
    "certifications": [
        {
            "name": "GRADE C ( S.O.B : 060 5126)",
            "date": "1999",
            "issuer": "MMABATHO SECURITY COLLEGE"
        }
    ],
    "projects": [],
    "leadership": {
        "positions": "VOLUNTEER AS ABET TEACHER",
        "volunteer": "1 year",
        "mentoring": "1 year"
    },
    "awards": {},
    "additional_info": {},
    "references": [
        {
            "name": "Mrs. Muriel Kealeboga Moloko",
            "relationship": "Educator",
            "contact": "0813845729"
        },
        {
            "name": "Mrs. S. G. Motsatsi",
            "relationship": "Ex-Educator",
            "contact": "0836980745"
        },
        {
            "name": "Mrs. L. P. PETLELE",
            "relationship": "Principal",
            "contact": "0822578655"
        }
    ]
}

# Display functions for Streamlit preview (same as before)
def display_personal_info(info, style):
    if style == "Traditional":
        st.header(f"{info['full_name']} {info['surname']}")
        st.write(f"**Location:** {info['p_location']}")
        st.write(f"**Phone:** {info['phone']}")
        st.write(f"**Email:** {info['email']}")
        if info['linkedin']:
            st.write(f"**LinkedIn:** {info['linkedin']}")
    elif style == "Creative":
        st.markdown(f"<h1 style='color: #4CAF50;'>{info['full_name']} {info['surname']}</h1>", unsafe_allow_html=True)
        st.markdown(f"<p><strong>📍 Location:</strong> {info['p_location']}</p>", unsafe_allow_html=True)
        st.markdown(f"<p><strong>📞 Phone:</strong> {info['phone']}</p>", unsafe_allow_html=True)
        st.markdown(f"<p><strong>✉️ Email:</strong> {info['email']}</p>", unsafe_allow_html=True)
        if info['linkedin']:
            st.markdown(f"<p><strong>🔗 LinkedIn:</strong> {info['linkedin']}</p>", unsafe_allow_html=True)
    elif style == "Minimalist":
        st.markdown(f"### {info['full_name']} {info['surname']}")
        st.markdown(f"{info['p_location']} | {info['phone']} | {info['email']}")
    elif style == "Modern":
        st.markdown(f"<h1 style='border-bottom: 2px solid #333;'>{info['full_name']} {info['surname']}</h1>", unsafe_allow_html=True)
        col1, col2, col3 = st.columns(3)
        col1.write(f"**Location:** {info['p_location']}")
        col2.write(f"**Phone:** {info['phone']}")
        col3.write(f"**Email:** {info['email']}")
    elif style == "Corporate":
        st.markdown(f"<h1 style='color: #003087;'>{info['full_name']} {info['surname']}</h1>", unsafe_allow_html=True)
        st.markdown(f"<p><strong>Location:</strong> {info['p_location']} | <strong>Phone:</strong> {info['phone']} | <strong>Email:</strong> {info['email']}</p>", unsafe_allow_html=True)
    elif style == "Academic":
        st.markdown(f"<h2 style='font-family: serif;'>{info['full_name']} {info['surname']}</h2>", unsafe_allow_html=True)
        st.write(f"Address: {info['p_location']}")
        st.write(f"Contact: {info['phone']} | {info['email']}")
    elif style == "Tech-Focused":
        st.markdown(f"<h1 style='color: #00bcd4;'>{info['full_name']} {info['surname']}</h1>", unsafe_allow_html=True)
        st.markdown(f"<p style='font-family: monospace;'>Location: {info['p_location']} | Contact: {info['phone']} | Email: {info['email']}</p>", unsafe_allow_html=True)

def display_work_experience(experience, style):
    if style == "Traditional":
        st.header("Work Experience")
        for job in experience:
            st.subheader(job['job_title'])
            st.write(f"**Company:** {job['company']}")
            st.write(f"**Dates:** {job['dates']}")
            st.write(f"**Reason for leaving:** {job['reason']}")
            st.write("---")
    elif style == "Creative":
        st.markdown("<h2 style='color: #2196F3;'>💼 Work Experience</h2>", unsafe_allow_html=True)
        for job in experience:
            st.markdown(f"<h3>{job['job_title']}</h3>", unsafe_allow_html=True)
            st.markdown(f"<p><strong>Company:</strong> {job['company']}</p>", unsafe_allow_html=True)
            st.markdown(f"<p><strong>Dates:</strong> {job['dates']}</p>", unsafe_allow_html=True)
            st.markdown(f"<p><strong>Reason for leaving:</strong> {job['reason']}</p>", unsafe_allow_html=True)
            st.markdown("<hr>", unsafe_allow_html=True)
    elif style == "Minimalist":
        st.markdown("### Work Experience")
        for job in experience:
            st.markdown(f"**{job['job_title']}** - {job['company']} ({job['dates']})")
            st.markdown(f"Reason: {job['reason']}")
    elif style == "Modern":
        st.markdown("<h2 style='border-bottom: 1px solid #ccc;'>Work Experience</h2>", unsafe_allow_html=True)
        for job in experience:
            st.markdown(f"<h3 style='color: #555;'>{job['job_title']}</h3>", unsafe_allow_html=True)
            st.markdown(f"{job['company']} | {job['dates']}")
            st.markdown(f"Reason for leaving: {job['reason']}")
            st.markdown("<hr style='border-top: 1px dashed #ccc;'>", unsafe_allow_html=True)
    elif style == "Corporate":
        st.markdown("<h2 style='color: #003087;'>Professional Experience</h2>", unsafe_allow_html=True)
        for job in experience:
            st.markdown(f"<h3>{job['job_title']}</h3>", unsafe_allow_html=True)
            st.markdown(f"<strong>{job['company']}</strong> | {job['dates']}")
            st.markdown(f"Reason for Departure: {job['reason']}")
            st.markdown("<hr>", unsafe_allow_html=True)
    elif style == "Academic":
        st.markdown("<h2 style='font-family: serif;'>Professional Appointments</h2>", unsafe_allow_html=True)
        for job in experience:
            st.markdown(f"**{job['job_title']}**, {job['company']}, {job['dates']}")
            st.markdown(f"Reason for leaving: {job['reason']}")
            st.markdown("---")
    elif style == "Tech-Focused":
        st.markdown("<h2 style='color: #00bcd4;'>[Work Experience]</h2>", unsafe_allow_html=True)
        for job in experience:
            st.markdown(f"<h3 style='font-family: monospace;'>{job['job_title']}</h3>", unsafe_allow_html=True)
            st.markdown(f"// {job['company']} ({job['dates']})")
            st.markdown(f"// Reason: {job['reason']}")
            st.markdown("<hr style='border-top: 1px dotted #00bcd4;'>", unsafe_allow_html=True)

def display_education(education, style):
    if style == "Traditional":
        st.header("Education")
        for edu in education:
            st.subheader(edu['degree'])
            st.write(f"**Field:** {edu['field']}")
            st.write(f"**Institution:** {edu['institution']}, {edu['location']}")
            st.write(f"**Graduation Date:** {edu['graduation_date']}")
            st.write("---")
    elif style == "Creative":
        st.markdown("<h2 style='color: #3F51B5;'>🎓 Education</h2>", unsafe_allow_html=True)
        for edu in education:
            st.markdown(f"<h3>{edu['degree']}</h3>", unsafe_allow_html=True)
            st.markdown(f"<p><strong>Field:</strong> {edu['field']}</p>", unsafe_allow_html=True)
            st.markdown(f"<p><strong>Institution:</strong> {edu['institution']}, {edu['location']}</p>", unsafe_allow_html=True)
            st.markdown(f"<p><strong>Graduation Date:</strong> {edu['graduation_date']}</p>", unsafe_allow_html=True)
            st.markdown("<hr>", unsafe_allow_html=True)
    elif style == "Minimalist":
        st.markdown("### Education")
        for edu in education:
            st.markdown(f"**{edu['degree']}** - {edu['institution']}, {edu['location']} ({edu['graduation_date']})")
    elif style == "Modern":
        st.markdown("<h2 style='border-bottom: 1px solid #ccc;'>Education</h2>", unsafe_allow_html=True)
        for edu in education:
            st.markdown(f"<h3 style='color: #555;'>{edu['degree']}</h3>", unsafe_allow_html=True)
            st.markdown(f"{edu['field']} | {edu['institution']}, {edu['location']}")
            st.markdown(f"Graduated: {edu['graduation_date']}")
            st.markdown("<hr style='border-top: 1px dashed #ccc;'>", unsafe_allow_html=True)
    elif style == "Corporate":
        st.markdown("<h2 style='color: #003087;'>Education</h2>", unsafe_allow_html=True)
        for edu in education:
            st.markdown(f"<h3>{edu['degree']}</h3>", unsafe_allow_html=True)
            st.markdown(f"{edu['field']}, {edu['institution']}, {edu['location']}")
            st.markdown(f"Graduation Year: {edu['graduation_date']}")
            st.markdown("<hr>", unsafe_allow_html=True)
    elif style == "Academic":
        st.markdown("<h2 style='font-family: serif;'>Academic Qualifications</h2>", unsafe_allow_html=True)
        for edu in education:
            st.markdown(f"**{edu['degree']}** in {edu['field']}, {edu['institution']}, {edu['location']}, {edu['graduation_date']}")
            st.markdown("---")
    elif style == "Tech-Focused":
        st.markdown("<h2 style='color: #00bcd4;'>[Education]</h2>", unsafe_allow_html=True)
        for edu in education:
            st.markdown(f"<h3 style='font-family: monospace;'>{edu['degree']}</h3>", unsafe_allow_html=True)
            st.markdown(f"// {edu['field']}, {edu['institution']}, {edu['location']}")
            st.markdown(f"// Completed: {edu['graduation_date']}")
            st.markdown("<hr style='border-top: 1px dotted #00bcd4;'>", unsafe_allow_html=True)

def display_skills(skills, style):
    if style == "Traditional":
        st.header("Skills")
        if skills["technical"]:
            st.subheader("Technical Skills")
            for skill in skills["technical"].split("\n"):
                if skill.strip():
                    st.markdown(f"- {skill}")
        if skills["software"]:
            st.subheader("Software")
            st.write(skills["software"])
        if skills["languages"]:
            st.subheader("Languages")
            for lang in skills["languages"].split("\n"):
                if lang.strip():
                    st.markdown(f"- {lang}")
        if skills["soft_skills"]:
            st.subheader("Soft Skills")
            for skill in skills["soft_skills"].split("\n"):
                if skill.strip():
                    st.markdown(f"- {skill}")
    elif style == "Creative":
        st.markdown("<h2 style='color: #FF9800;'>🛠️ Skills</h2>", unsafe_allow_html=True)
        if skills["technical"]:
            st.markdown("<h3>Technical Skills</h3>", unsafe_allow_html=True)
            for skill in skills["technical"].split("\n"):
                if skill.strip():
                    st.markdown(f"- {skill}", unsafe_allow_html=True)
        if skills["software"]:
            st.markdown("<h3>Software</h3>", unsafe_allow_html=True)
            st.markdown(f"<p>{skills['software']}</p>", unsafe_allow_html=True)
        if skills["languages"]:
            st.markdown("<h3>Languages</h3>", unsafe_allow_html=True)
            for lang in skills["languages"].split("\n"):
                if lang.strip():
                    st.markdown(f"- {lang}", unsafe_allow_html=True)
        if skills["soft_skills"]:
            st.markdown("<h3>Soft Skills</h3>", unsafe_allow_html=True)
            for skill in skills["soft_skills"].split("\n"):
                if skill.strip():
                    st.markdown(f"- {skill}", unsafe_allow_html=True)
    elif style == "Minimalist":
        st.markdown("### Skills")
        if skills["technical"]:
            st.markdown("**Technical:** " + ", ".join([s.strip() for s in skills["technical"].split("\n") if s.strip()]))
        if skills["software"]:
            st.markdown(f"**Software:** {skills['software']}")
        if skills["languages"]:
            st.markdown("**Languages:** " + ", ".join([s.strip() for s in skills["languages"].split("\n") if s.strip()]))
        if skills["soft_skills"]:
            st.markdown("**Soft Skills:** " + ", ".join([s.strip() for s in skills["soft_skills"].split("\n") if s.strip()]))
    elif style == "Modern":
        st.markdown("<h2 style='border-bottom: 1px solid #ccc;'>Skills</h2>", unsafe_allow_html=True)
        if skills["technical"]:
            st.markdown("<h3 style='color: #555;'>Technical Skills</h3>", unsafe_allow_html=True)
            st.markdown(", ".join([s.strip() for s in skills["technical"].split("\n") if s.strip()]))
        if skills["software"]:
            st.markdown("<h3 style='color: #555;'>Software</h3>", unsafe_allow_html=True)
            st.markdown(skills["software"])
        if skills["languages"]:
            st.markdown("<h3 style='color: #555;'>Languages</h3>", unsafe_allow_html=True)
            st.markdown(", ".join([s.strip() for s in skills["languages"].split("\n") if s.strip()]))
        if skills["soft_skills"]:
            st.markdown("<h3 style='color: #555;'>Soft Skills</h3>", unsafe_allow_html=True)
            st.markdown(", ".join([s.strip() for s in skills["soft_skills"].split("\n") if s.strip()]))
    elif style == "Corporate":
        st.markdown("<h2 style='color: #003087;'>Core Competencies</h2>", unsafe_allow_html=True)
        if skills["technical"]:
            st.markdown("**Technical Skills:** " + ", ".join([s.strip() for s in skills["technical"].split("\n") if s.strip()]))
        if skills["software"]:
            st.markdown(f"**Software Proficiency:** {skills['software']}")
        if skills["languages"]:
            st.markdown("**Languages:** " + ", ".join([s.strip() for s in skills["languages"].split("\n") if s.strip()]))
        if skills["soft_skills"]:
            st.markdown("**Professional Skills:** " + ", ".join([s.strip() for s in skills["soft_skills"].split("\n") if s.strip()]))
    elif style == "Academic":
        st.markdown("<h2 style='font-family: serif;'>Skills and Proficiencies</h2>", unsafe_allow_html=True)
        if skills["technical"]:
            st.markdown("**Technical Skills:** " + "; ".join([s.strip() for s in skills["technical"].split("\n") if s.strip()]))
        if skills["software"]:
            st.markdown(f"**Software Knowledge:** {skills['software']}")
        if skills["languages"]:
            st.markdown("**Languages:** " + "; ".join([s.strip() for s in skills["languages"].split("\n") if s.strip()]))
        if skills["soft_skills"]:
            st.markdown("**Interpersonal Skills:** " + "; ".join([s.strip() for s in skills["soft_skills"].split("\n") if s.strip()]))
    elif style == "Tech-Focused":
        st.markdown("<h2 style='color: #00bcd4;'>[Skills]</h2>", unsafe_allow_html=True)
        if skills["technical"]:
            st.markdown("<h3 style='font-family: monospace;'>// Technical</h3>", unsafe_allow_html=True)
            for skill in skills["technical"].split("\n"):
                if skill.strip():
                    st.markdown(f"> {skill}", unsafe_allow_html=True)
        if skills["software"]:
            st.markdown("<h3 style='font-family: monospace;'>// Software</h3>", unsafe_allow_html=True)
            st.markdown(f"> {skills['software']}")
        if skills["languages"]:
            st.markdown("<h3 style='font-family: monospace;'>// Languages</h3>", unsafe_allow_html=True)
            for lang in skills["languages"].split("\n"):
                if lang.strip():
                    st.markdown(f"> {lang}", unsafe_allow_html=True)
        if skills["soft_skills"]:
            st.markdown("<h3 style='font-family: monospace;'>// Soft Skills</h3>", unsafe_allow_html=True)
            for skill in skills["soft_skills"].split("\n"):
                if skill.strip():
                    st.markdown(f"> {skill}", unsafe_allow_html=True)

def display_certifications(certifications, style):
    if style == "Traditional":
        st.header("Certifications")
        for cert in certifications:
            st.write(f"**{cert['name']}** - {cert['issuer']} ({cert['date']})")
    elif style == "Creative":
        st.markdown("<h2 style='color: #9C27B0;'>🏅 Certifications</h2>", unsafe_allow_html=True)
        for cert in certifications:
            st.markdown(f"<p><strong>{cert['name']}</strong> - {cert['issuer']} ({cert['date']})</p>", unsafe_allow_html=True)
    elif style == "Minimalist":
        st.markdown("### Certifications")
        for cert in certifications:
            st.markdown(f"{cert['name']} - {cert['issuer']} ({cert['date']})")
    elif style == "Modern":
        st.markdown("<h2 style='border-bottom: 1px solid #ccc;'>Certifications</h2>", unsafe_allow_html=True)
        for cert in certifications:
            st.markdown(f"<p style='color: #555;'>{cert['name']} - {cert['issuer']} ({cert['date']})</p>", unsafe_allow_html=True)
    elif style == "Corporate":
        st.markdown("<h2 style='color: #003087;'>Certifications</h2>", unsafe_allow_html=True)
        for cert in certifications:
            st.markdown(f"{cert['name']} - {cert['issuer']}, {cert['date']}")
    elif style == "Academic":
        st.markdown("<h2 style='font-family: serif;'>Certifications and Training</h2>", unsafe_allow_html=True)
        for cert in certifications:
            st.markdown(f"**{cert['name']}**, {cert['issuer']}, {cert['date']}")
    elif style == "Tech-Focused":
        st.markdown("<h2 style='color: #00bcd4;'>[Certifications]</h2>", unsafe_allow_html=True)
        for cert in certifications:
            st.markdown(f"<p style='font-family: monospace;'>// {cert['name']} - {cert['issuer']} ({cert['date']})</p>", unsafe_allow_html=True)

def display_leadership(leadership, style):
    if style == "Traditional":
        st.header("Leadership and Volunteering")
        if leadership["positions"]:
            st.write(f"- {leadership['positions']}: {leadership['volunteer']}")
        if leadership["mentoring"]:
            st.write(f"- Mentoring: {leadership['mentoring']}")
    elif style == "Creative":
        st.markdown("<h2 style='color: #FFC107;'>🤝 Leadership and Volunteering</h2>", unsafe_allow_html=True)
        if leadership["positions"]:
            st.markdown(f"<p>- {leadership['positions']}: {leadership['volunteer']}</p>", unsafe_allow_html=True)
        if leadership["mentoring"]:
            st.markdown(f"<p>- Mentoring: {leadership['mentoring']}</p>", unsafe_allow_html=True)
    elif style == "Minimalist":
        st.markdown("### Leadership and Volunteering")
        if leadership["positions"]:
            st.markdown(f"{leadership['positions']}: {leadership['volunteer']}")
        if leadership["mentoring"]:
            st.markdown(f"Mentoring: {leadership['mentoring']}")
    elif style == "Modern":
        st.markdown("<h2 style='border-bottom: 1px solid #ccc;'>Leadership and Volunteering</h2>", unsafe_allow_html=True)
        if leadership["positions"]:
            st.markdown(f"<p style='color: #555;'>{leadership['positions']}: {leadership['volunteer']}</p>", unsafe_allow_html=True)
        if leadership["mentoring"]:
            st.markdown(f"<p style='color: #555;'>Mentoring: {leadership['mentoring']}</p>", unsafe_allow_html=True)
    elif style == "Corporate":
        st.markdown("<h2 style='color: #003087;'>Leadership and Community Involvement</h2>", unsafe_allow_html=True)
        if leadership["positions"]:
            st.markdown(f"{leadership['positions']}: {leadership['volunteer']}")
        if leadership["mentoring"]:
            st.markdown(f"Mentoring Experience: {leadership['mentoring']}")
    elif style == "Academic":
        st.markdown("<h2 style='font-family: serif;'>Leadership and Service</h2>", unsafe_allow_html=True)
        if leadership["positions"]:
            st.markdown(f"**{leadership['positions']}**, Duration: {leadership['volunteer']}")
        if leadership["mentoring"]:
            st.markdown(f"**Mentoring**, Duration: {leadership['mentoring']}")
    elif style == "Tech-Focused":
        st.markdown("<h2 style='color: #00bcd4;'>[Leadership]</h2>", unsafe_allow_html=True)
        if leadership["positions"]:
            st.markdown(f"<p style='font-family: monospace;'>// {leadership['positions']}: {leadership['volunteer']}</p>", unsafe_allow_html=True)
        if leadership["mentoring"]:
            st.markdown(f"<p style='font-family: monospace;'>// Mentoring: {leadership['mentoring']}</p>", unsafe_allow_html=True)

def display_references(references, style):
    if style == "Traditional":
        st.header("References")
        for ref in references:
            st.write(f"**{ref['name']}** - {ref['relationship']} - {ref['contact']}")
    elif style == "Creative":
        st.markdown("<h2 style='color: #E91E63;'>📞 References</h2>", unsafe_allow_html=True)
        for ref in references:
            st.markdown(f"<p><strong>{ref['name']}</strong> - {ref['relationship']} - {ref['contact']}</p>", unsafe_allow_html=True)
    elif style == "Minimalist":
        st.markdown("### References")
        for ref in references:
            st.markdown(f"{ref['name']} - {ref['relationship']} - {ref['contact']}")
    elif style == "Modern":
        st.markdown("<h2 style='border-bottom: 1px solid #ccc;'>References</h2>", unsafe_allow_html=True)
        for ref in references:
            st.markdown(f"<p style='color: #555;'>{ref['name']} - {ref['relationship']} - {ref['contact']}</p>", unsafe_allow_html=True)
    elif style == "Corporate":
        st.markdown("<h2 style='color: #003087;'>References</h2>", unsafe_allow_html=True)
        for ref in references:
            st.markdown(f"{ref['name']}, {ref['relationship']}, Contact: {ref['contact']}")
    elif style == "Academic":
        st.markdown("<h2 style='font-family: serif;'>References</h2>", unsafe_allow_html=True)
        for ref in references:
            st.markdown(f"**{ref['name']}**, {ref['relationship']}, Contact: {ref['contact']}")
    elif style == "Tech-Focused":
        st.markdown("<h2 style='color: #00bcd4;'>[References]</h2>", unsafe_allow_html=True)
        for ref in references:
            st.markdown(f"<p style='font-family: monospace;'>// {ref['name']} - {ref['relationship']} - {ref['contact']}</p>", unsafe_allow_html=True)

# Functions to generate Word document
def add_heading(doc, text, level, style, color=None):
    heading = doc.add_heading(text, level=level)
    heading.style.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)
    if color:
        heading.style.font.color.rgb = color
    if style == "Tech-Focused" and level == 2:
        heading.style.font.name = "Courier New"

def add_paragraph(doc, text, bold=False, color=None, font_name=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    if font_name:
        run.font.name = font_name
    run.font.size = Pt(11)

def generate_docx(data, style):
    doc = Document()

    # Define styles for headings
    styles = doc.styles
    if "Heading 1" not in styles:
        styles.add_style("Heading 1", WD_STYLE_TYPE.PARAGRAPH)
    if "Heading 2" not in styles:
        styles.add_style("Heading 2", WD_STYLE_TYPE.PARAGRAPH)

    # Personal Info
    info = data["personal_info"]
    if style == "Traditional":
        add_heading(doc, f"{info['full_name']} {info['surname']}", 1, style)
        add_paragraph(doc, f"Location: {info['p_location']}")
        add_paragraph(doc, f"Phone: {info['phone']}")
        add_paragraph(doc, f"Email: {info['email']}")
    elif style == "Creative":
        add_heading(doc, f"{info['full_name']} {info['surname']}", 1, style, RGBColor(76, 175, 80))
        add_paragraph(doc, f"📍 Location: {info['p_location']}")
        add_paragraph(doc, f"📞 Phone: {info['phone']}")
        add_paragraph(doc, f"✉️ Email: {info['email']}")
    elif style == "Minimalist":
        add_heading(doc, f"{info['full_name']} {info['surname']}", 1, style)
        add_paragraph(doc, f"{info['p_location']} | {info['phone']} | {info['email']}")
    elif style == "Modern":
        add_heading(doc, f"{info['full_name']} {info['surname']}", 1, style)
        add_paragraph(doc, f"Location: {info['p_location']} | Phone: {info['phone']} | Email: {info['email']}")
    elif style == "Corporate":
        add_heading(doc, f"{info['full_name']} {info['surname']}", 1, style, RGBColor(0, 48, 135))
        add_paragraph(doc, f"Location: {info['p_location']} | Phone: {info['phone']} | Email: {info['email']}")
    elif style == "Academic":
        add_heading(doc, f"{info['full_name']} {info['surname']}", 1, style)
        add_paragraph(doc, f"Address: {info['p_location']}")
        add_paragraph(doc, f"Contact: {info['phone']} | {info['email']}")
    elif style == "Tech-Focused":
        add_heading(doc, f"{info['full_name']} {info['surname']}", 1, style, RGBColor(0, 188, 212))
        add_paragraph(doc, f"Location: {info['p_location']} | Contact: {info['phone']} | Email: {info['email']}", font_name="Courier New")

    # Work Experience
    experience = data["work_experience"]
    if style == "Traditional":
        add_heading(doc, "Work Experience", 2, style)
        for job in experience:
            add_paragraph(doc, job['job_title'], bold=True)
            add_paragraph(doc, f"Company: {job['company']}")
            add_paragraph(doc, f"Dates: {job['dates']}")
            add_paragraph(doc, f"Reason for leaving: {job['reason']}")
            doc.add_paragraph("---")
    elif style == "Creative":
        add_heading(doc, "💼 Work Experience", 2, style, RGBColor(33, 150, 243))
        for job in experience:
            add_paragraph(doc, job['job_title'], bold=True)
            add_paragraph(doc, f"Company: {job['company']}")
            add_paragraph(doc, f"Dates: {job['dates']}")
            add_paragraph(doc, f"Reason for leaving: {job['reason']}")
            doc.add_paragraph("---")
    elif style == "Minimalist":
        add_heading(doc, "Work Experience", 2, style)
        for job in experience:
            add_paragraph(doc, f"{job['job_title']} - {job['company']} ({job['dates']})", bold=True)
            add_paragraph(doc, f"Reason: {job['reason']}")
    elif style == "Modern":
        add_heading(doc, "Work Experience", 2, style)
        for job in experience:
            add_paragraph(doc, job['job_title'], bold=True)
            add_paragraph(doc, f"{job['company']} | {job['dates']}")
            add_paragraph(doc, f"Reason for leaving: {job['reason']}")
            doc.add_paragraph("---")
    elif style == "Corporate":
        add_heading(doc, "Professional Experience", 2, style, RGBColor(0, 48, 135))
        for job in experience:
            add_paragraph(doc, job['job_title'], bold=True)
            add_paragraph(doc, f"{job['company']} | {job['dates']}")
            add_paragraph(doc, f"Reason for Departure: {job['reason']}")
            doc.add_paragraph("---")
    elif style == "Academic":
        add_heading(doc, "Professional Appointments", 2, style)
        for job in experience:
            add_paragraph(doc, f"{job['job_title']}, {job['company']}, {job['dates']}", bold=True)
            add_paragraph(doc, f"Reason for leaving: {job['reason']}")
            doc.add_paragraph("---")
    elif style == "Tech-Focused":
        add_heading(doc, "[Work Experience]", 2, style, RGBColor(0, 188, 212))
        for job in experience:
            add_paragraph(doc, job['job_title'], bold=True, font_name="Courier New")
            add_paragraph(doc, f"// {job['company']} ({job['dates']})", font_name="Courier New")
            add_paragraph(doc, f"// Reason: {job['reason']}", font_name="Courier New")
            doc.add_paragraph("---")

    # Education
    education = data["education"]
    if style == "Traditional":
        add_heading(doc, "Education", 2, style)
        for edu in education:
            add_paragraph(doc, edu['degree'], bold=True)
            add_paragraph(doc, f"Field: {edu['field']}")
            add_paragraph(doc, f"Institution: {edu['institution']}, {edu['location']}")
            add_paragraph(doc, f"Graduation Date: {edu['graduation_date']}")
            doc.add_paragraph("---")
    elif style == "Creative":
        add_heading(doc, "🎓 Education", 2, style, RGBColor(63, 81, 181))
        for edu in education:
            add_paragraph(doc, edu['degree'], bold=True)
            add_paragraph(doc, f"Field: {edu['field']}")
            add_paragraph(doc, f"Institution: {edu['institution']}, {edu['location']}")
            add_paragraph(doc, f"Graduation Date: {edu['graduation_date']}")
            doc.add_paragraph("---")
    elif style == "Minimalist":
        add_heading(doc, "Education", 2, style)
        for edu in education:
            add_paragraph(doc, f"{edu['degree']} - {edu['institution']}, {edu['location']} ({edu['graduation_date']})", bold=True)
    elif style == "Modern":
        add_heading(doc, "Education", 2, style)
        for edu in education:
            add_paragraph(doc, edu['degree'], bold=True)
            add_paragraph(doc, f"{edu['field']} | {edu['institution']}, {edu['location']}")
            add_paragraph(doc, f"Graduated: {edu['graduation_date']}")
            doc.add_paragraph("---")
    elif style == "Corporate":
        add_heading(doc, "Education", 2, style, RGBColor(0, 48, 135))
        for edu in education:
            add_paragraph(doc, edu['degree'], bold=True)
            add_paragraph(doc, f"{edu['field']}, {edu['institution']}, {edu['location']}")
            add_paragraph(doc, f"Graduation Year: {edu['graduation_date']}")
            doc.add_paragraph("---")
    elif style == "Academic":
        add_heading(doc, "Academic Qualifications", 2, style)
        for edu in education:
            add_paragraph(doc, f"{edu['degree']} in {edu['field']}, {edu['institution']}, {edu['location']}, {edu['graduation_date']}", bold=True)
            doc.add_paragraph("---")
    elif style == "Tech-Focused":
        add_heading(doc, "[Education]", 2, style, RGBColor(0, 188, 212))
        for edu in education:
            add_paragraph(doc, edu['degree'], bold=True, font_name="Courier New")
            add_paragraph(doc, f"// {edu['field']}, {edu['institution']}, {edu['location']}", font_name="Courier New")
            add_paragraph(doc, f"// Completed: {edu['graduation_date']}", font_name="Courier New")
            doc.add_paragraph("---")

    # Skills
    skills = data["skills"]
    if style == "Traditional":
        add_heading(doc, "Skills", 2, style)
        if skills["technical"]:
            add_paragraph(doc, "Technical Skills", bold=True)
            for skill in skills["technical"].split("\n"):
                if skill.strip():
                    doc.add_paragraph(f"- {skill}", style="List Bullet")
        if skills["software"]:
            add_paragraph(doc, "Software", bold=True)
            add_paragraph(doc, skills["software"])
        if skills["languages"]:
            add_paragraph(doc, "Languages", bold=True)
            for lang in skills["languages"].split("\n"):
                if lang.strip():
                    doc.add_paragraph(f"- {lang}", style="List Bullet")
        if skills["soft_skills"]:
            add_paragraph(doc, "Soft Skills", bold=True)
            for skill in skills["soft_skills"].split("\n"):
                if skill.strip():
                    doc.add_paragraph(f"- {skill}", style="List Bullet")
    elif style == "Creative":
        add_heading(doc, "🛠️ Skills", 2, style, RGBColor(255, 152, 0))
        if skills["technical"]:
            add_paragraph(doc, "Technical Skills", bold=True)
            for skill in skills["technical"].split("\n"):
                if skill.strip():
                    doc.add_paragraph(f"- {skill}", style="List Bullet")
        if skills["software"]:
            add_paragraph(doc, "Software", bold=True)
            add_paragraph(doc, skills["software"])
        if skills["languages"]:
            add_paragraph(doc, "Languages", bold=True)
            for lang in skills["languages"].split("\n"):
                if lang.strip():
                    doc.add_paragraph(f"- {lang}", style="List Bullet")
        if skills["soft_skills"]:
            add_paragraph(doc, "Soft Skills", bold=True)
            for skill in skills["soft_skills"].split("\n"):
                if skill.strip():
                    doc.add_paragraph(f"- {skill}", style="List Bullet")
    elif style == "Minimalist":
        add_heading(doc, "Skills", 2, style)
        if skills["technical"]:
            add_paragraph(doc, f"Technical: {', '.join([s.strip() for s in skills['technical'].split('\n') if s.strip()])}")
        if skills["software"]:
            add_paragraph(doc, f"Software: {skills['software']}")
        if skills["languages"]:
            add_paragraph(doc, f"Languages: {', '.join([s.strip() for s in skills['languages'].split('\n') if s.strip()])}")
        if skills["soft_skills"]:
            add_paragraph(doc, f"Soft Skills: {', '.join([s.strip() for s in skills['soft_skills'].split('\n') if s.strip()])}")
    elif style == "Modern":
        add_heading(doc, "Skills", 2, style)
        if skills["technical"]:
            add_paragraph(doc, "Technical Skills", bold=True)
            add_paragraph(doc, ", ".join([s.strip() for s in skills["technical"].split("\n") if s.strip()]))
        if skills["software"]:
            add_paragraph(doc, "Software", bold=True)
            add_paragraph(doc, skills["software"])
        if skills["languages"]:
            add_paragraph(doc, "Languages", bold=True)
            add_paragraph(doc, ", ".join([s.strip() for s in skills["languages"].split("\n") if s.strip()]))
        if skills["soft_skills"]:
            add_paragraph(doc, "Soft Skills", bold=True)
            add_paragraph(doc, ", ".join([s.strip() for s in skills["soft_skills"].split("\n") if s.strip()]))
    elif style == "Corporate":
        add_heading(doc, "Core Competencies", 2, style, RGBColor(0, 48, 135))
        if skills["technical"]:
            add_paragraph(doc, f"Technical Skills: {', '.join([s.strip() for s in skills['technical'].split('\n') if s.strip()])}")
        if skills["software"]:
            add_paragraph(doc, f"Software Proficiency: {skills['software']}")
        if skills["languages"]:
            add_paragraph(doc, f"Languages: {', '.join([s.strip() for s in skills['languages'].split('\n') if s.strip()])}")
        if skills["soft_skills"]:
            add_paragraph(doc, f"Professional Skills: {', '.join([s.strip() for s in skills['soft_skills'].split('\n') if s.strip()])}")
    elif style == "Academic":
        add_heading(doc, "Skills and Proficiencies", 2, style)
        if skills["technical"]:
            add_paragraph(doc, f"Technical Skills: {'; '.join([s.strip() for s in skills['technical'].split('\n') if s.strip()])}")
        if skills["software"]:
            add_paragraph(doc, f"Software Knowledge: {skills['software']}")
        if skills["languages"]:
            add_paragraph(doc, f"Languages: {'; '.join([s.strip() for s in skills['languages'].split('\n') if s.strip()])}")
        if skills["soft_skills"]:
            add_paragraph(doc, f"Interpersonal Skills: {'; '.join([s.strip() for s in skills['soft_skills'].split('\n') if s.strip()])}")
    elif style == "Tech-Focused":
        add_heading(doc, "[Skills]", 2, style, RGBColor(0, 188, 212))
        if skills["technical"]:
            add_paragraph(doc, "// Technical", bold=True, font_name="Courier New")
            for skill in skills["technical"].split("\n"):
                if skill.strip():
                    add_paragraph(doc, f"> {skill}", font_name="Courier New")
        if skills["software"]:
            add_paragraph(doc, "// Software", bold=True, font_name="Courier New")
            add_paragraph(doc, f"> {skills['software']}", font_name="Courier New")
        if skills["languages"]:
            add_paragraph(doc, "// Languages", bold=True, font_name="Courier New")
            for lang in skills["languages"].split("\n"):
                if lang.strip():
                    add_paragraph(doc, f"> {lang}", font_name="Courier New")
        if skills["soft_skills"]:
            add_paragraph(doc, "// Soft Skills", bold=True, font_name="Courier New")
            for skill in skills["soft_skills"].split("\n"):
                if skill.strip():
                    add_paragraph(doc, f"> {skill}", font_name="Courier New")

    # Certifications
    certifications = data["certifications"]
    if style == "Traditional":
        add_heading(doc, "Certifications", 2, style)
        for cert in certifications:
            add_paragraph(doc, f"{cert['name']} - {cert['issuer']} ({cert['date']})", bold=True)
    elif style == "Creative":
        add_heading(doc, "🏅 Certifications", 2, style, RGBColor(156, 39, 176))
        for cert in certifications:
            add_paragraph(doc, f"{cert['name']} - {cert['issuer']} ({cert['date']})")
    elif style == "Minimalist":
        add_heading(doc, "Certifications", 2, style)
        for cert in certifications:
            add_paragraph(doc, f"{cert['name']} - {cert['issuer']} ({cert['date']})")
    elif style == "Modern":
        add_heading(doc, "Certifications", 2, style)
        for cert in certifications:
            add_paragraph(doc, f"{cert['name']} - {cert['issuer']} ({cert['date']})")
    elif style == "Corporate":
        add_heading(doc, "Certifications", 2, style, RGBColor(0, 48, 135))
        for cert in certifications:
            add_paragraph(doc, f"{cert['name']} - {cert['issuer']}, {cert['date']}")
    elif style == "Academic":
        add_heading(doc, "Certifications and Training", 2, style)
        for cert in certifications:
            add_paragraph(doc, f"{cert['name']}, {cert['issuer']}, {cert['date']}", bold=True)
    elif style == "Tech-Focused":
        add_heading(doc, "[Certifications]", 2, style, RGBColor(0, 188, 212))
        for cert in certifications:
            add_paragraph(doc, f"// {cert['name']} - {cert['issuer']} ({cert['date']})", font_name="Courier New")

    # Leadership
    leadership = data["leadership"]
    if style == "Traditional":
        add_heading(doc, "Leadership and Volunteering", 2, style)
        if leadership["positions"]:
            add_paragraph(doc, f"- {leadership['positions']}: {leadership['volunteer']}")
        if leadership["mentoring"]:
            add_paragraph(doc, f"- Mentoring: {leadership['mentoring']}")
    elif style == "Creative":
        add_heading(doc, "🤝 Leadership and Volunteering", 2, style, RGBColor(255, 193, 7))
        if leadership["positions"]:
            add_paragraph(doc, f"- {leadership['positions']}: {leadership['volunteer']}")
        if leadership["mentoring"]:
            add_paragraph(doc, f"- Mentoring: {leadership['mentoring']}")
    elif style == "Minimalist":
        add_heading(doc, "Leadership and Volunteering", 2, style)
        if leadership["positions"]:
            add_paragraph(doc, f"{leadership['positions']}: {leadership['volunteer']}")
        if leadership["mentoring"]:
            add_paragraph(doc, f"Mentoring: {leadership['mentoring']}")
    elif style == "Modern":
        add_heading(doc, "Leadership and Volunteering", 2, style)
        if leadership["positions"]:
            add_paragraph(doc, f"{leadership['positions']}: {leadership['volunteer']}")
        if leadership["mentoring"]:
            add_paragraph(doc, f"Mentoring: {leadership['mentoring']}")
    elif style == "Corporate":
        add_heading(doc, "Leadership and Community Involvement", 2, style, RGBColor(0, 48, 135))
        if leadership["positions"]:
            add_paragraph(doc, f"{leadership['positions']}: {leadership['volunteer']}")
        if leadership["mentoring"]:
            add_paragraph(doc, f"Mentoring Experience: {leadership['mentoring']}")
    elif style == "Academic":
        add_heading(doc, "Leadership and Service", 2, style)
        if leadership["positions"]:
            add_paragraph(doc, f"{leadership['positions']}, Duration: {leadership['volunteer']}", bold=True)
        if leadership["mentoring"]:
            add_paragraph(doc, f"Mentoring, Duration: {leadership['mentoring']}", bold=True)
    elif style == "Tech-Focused":
        add_heading(doc, "[Leadership]", 2, style, RGBColor(0, 188, 212))
        if leadership["positions"]:
            add_paragraph(doc, f"// {leadership['positions']}: {leadership['volunteer']}", font_name="Courier New")
        if leadership["mentoring"]:
            add_paragraph(doc, f"// Mentoring: {leadership['mentoring']}", font_name="Courier New")

    # References
    references = data["references"]
    if style == "Traditional":
        add_heading(doc, "References", 2, style)
        for ref in references:
            add_paragraph(doc, f"{ref['name']} - {ref['relationship']} - {ref['contact']}", bold=True)
    elif style == "Creative":
        add_heading(doc, "📞 References", 2, style, RGBColor(233, 30, 99))
        for ref in references:
            add_paragraph(doc, f"{ref['name']} - {ref['relationship']} - {ref['contact']}")
    elif style == "Minimalist":
        add_heading(doc, "References", 2, style)
        for ref in references:
            add_paragraph(doc, f"{ref['name']} - {ref['relationship']} - {ref['contact']}")
    elif style == "Modern":
        add_heading(doc, "References", 2, style)
        for ref in references:
            add_paragraph(doc, f"{ref['name']} - {ref['relationship']} - {ref['contact']}")
    elif style == "Corporate":
        add_heading(doc, "References", 2, style, RGBColor(0, 48, 135))
        for ref in references:
            add_paragraph(doc, f"{ref['name']}, {ref['relationship']}, Contact: {ref['contact']}")
    elif style == "Academic":
        add_heading(doc, "References", 2, style)
        for ref in references:
            add_paragraph(doc, f"{ref['name']}, {ref['relationship']}, Contact: {ref['contact']}", bold=True)
    elif style == "Tech-Focused":
        add_heading(doc, "[References]", 2, style, RGBColor(0, 188, 212))
        for ref in references:
            add_paragraph(doc, f"// {ref['name']} - {ref['relationship']} - {ref['contact']}", font_name="Courier New")

    # Save document to a BytesIO stream
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Main Streamlit app
st.title("Resume Preview")

# Sidebar for template style selection
style = st.pills("Select Template Style", [
    "Traditional", "Creative", "Minimalist", "Modern", "Corporate", "Academic", "Tech-Focused"
])

# Display resume sections based on selected style
display_personal_info(resume_data["personal_info"], style)
display_work_experience(resume_data["work_experience"], style)
display_education(resume_data["education"], style)
display_skills(resume_data["skills"], style)
display_certifications(resume_data["certifications"], style)
display_leadership(resume_data["leadership"], style)
display_references(resume_data["references"], style)

# Add Save as Word button
if st.button("Save as Word Document"):
    doc_buffer = generate_docx(resume_data, style)
    st.download_button(
        label="Download Resume.docx",
        data=doc_buffer,
        file_name="Resume.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )